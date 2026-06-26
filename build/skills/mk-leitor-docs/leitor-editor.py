#!/usr/bin/env python3
import sys, json
from pathlib import Path

try:
    import PyPDF2
    from openpyxl import load_workbook
    from docx import Document
except:
    print(json.dumps({"erro": "Instale: pip install PyPDF2 openpyxl python-docx"}))
    sys.exit(1)

# ===== LER =====
def ler_pdf(arquivo):
    with open(arquivo, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        return "\n".join(p.extract_text() for p in reader.pages)

def ler_xlsx(arquivo):
    wb = load_workbook(arquivo)
    dados = {}
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        dados[sheet] = [[cell.value for cell in row] for row in ws.iter_rows()]
    return dados

def ler_docx(arquivo):
    doc = Document(arquivo)
    return "\n".join(p.text for p in doc.paragraphs)

# ===== EDITAR =====
def editar_pdf(arquivo, pagina, texto_novo):
    with open(arquivo, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        writer = PyPDF2.PdfWriter()
        for i, page in enumerate(reader.pages):
            writer.add_page(page)
    with open(arquivo, 'wb') as f:
        writer.write(f)
    return {"status": "PDF salvo"}

def editar_xlsx(arquivo, sheet, linha, coluna, valor):
    wb = load_workbook(arquivo)
    ws = wb[sheet]
    ws.cell(row=linha, column=coluna, value=valor)
    wb.save(arquivo)
    return {"status": f"Célula [{sheet}!{linha},{coluna}] = {valor}"}

def editar_docx(arquivo, paragrafo_idx, novo_texto):
    doc = Document(arquivo)
    if paragrafo_idx < len(doc.paragraphs):
        doc.paragraphs[paragrafo_idx].text = novo_texto
    doc.save(arquivo)
    return {"status": f"Parágrafo {paragrafo_idx} atualizado"}

# ===== MAIN =====
if len(sys.argv) < 2:
    print(json.dumps({"erro": "Use: python leitor-editor.py <acao> <arquivo> [opcoes]"}))
    sys.exit(1)

acao = sys.argv[1]
arquivo = sys.argv[2] if len(sys.argv) > 2 else None
ext = Path(arquivo).suffix.lower() if arquivo else None

try:
    if acao == "ler":
        if ext == '.pdf':
            resultado = ler_pdf(arquivo)
        elif ext == '.xlsx':
            resultado = ler_xlsx(arquivo)
        elif ext == '.docx':
            resultado = ler_docx(arquivo)
        else:
            resultado = {"erro": f"Formato {ext} não suportado"}
        print(json.dumps({"sucesso": True, "dados": resultado}))
    
    elif acao == "editar-xlsx":
        sheet, linha, coluna, valor = sys.argv[3], int(sys.argv[4]), int(sys.argv[5]), sys.argv[6]
        resultado = editar_xlsx(arquivo, sheet, linha, coluna, valor)
        print(json.dumps({"sucesso": True, "resultado": resultado}))
    
    elif acao == "editar-docx":
        paragrafo_idx, novo_texto = int(sys.argv[3]), sys.argv[4]
        resultado = editar_docx(arquivo, paragrafo_idx, novo_texto)
        print(json.dumps({"sucesso": True, "resultado": resultado}))
    
    else:
        print(json.dumps({"erro": f"Ação '{acao}' desconhecida"}))

except Exception as e:
    print(json.dumps({"erro": str(e)}))
